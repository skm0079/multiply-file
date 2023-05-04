import uuid
from faker import Faker
import random
from typing import List
from PIL import Image
import pytesseract
import re
from docx import Document
from datetime import datetime
from docxtpl import DocxTemplate
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx2pdf import convert
import pandas as pd
import yaml
import pprint
import tempfile
import pdf2image
import os
from docx2pdf import convert
from PIL import Image, ImageDraw, ImageFont
import xml.etree.ElementTree as ET

def annotate_text(doc_file_path, target_annotation_pairs, output_png_path, output_xml_path):
    """
    Annotate target text with annotation text in bounding boxes and convert to PNG image.

    Args:
        doc_file_path (str): Path to the input Docx file.
        target_annotation_pairs (list): A list of tuples, where each tuple contains a target text and its corresponding annotation text.
        output_png_path (str): Path to the output PNG file.
        output_xml_path (str): Path to the output XML file.

    Returns:
        None
    """
    # Check if input file exists
    if not os.path.isfile(doc_file_path):
        raise ValueError("Input file does not exist.")

    # Convert docx to png
    try:
        image = convert_to_image(doc_file_path)
    except Exception as e:
        raise RuntimeError("Failed to convert Docx to PNG.") from e

    # Initialize list of bounding boxes for all target text
    all_bbox = []

    # Annotate image with bounding boxes
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype("arial.ttf", 20)
    for target_text, annotation_text in target_annotation_pairs:
        # Get bounding boxes for current target text
        bbox = get_bounding_box_multi_page(target_text, image)
        all_bbox.extend(bbox)

        # Draw bounding boxes and annotation text for current target text
        if input_yaml_data['mark_annotations']:
            for box in bbox:
                draw.rectangle(box, outline="red", width=2)
                draw.text((box[0], box[1] - 20), annotation_text, font=font, fill="red")

    # Save annotated image to PNG file
    try:
        image.save(output_png_path, "PNG")
    except Exception as e:
        raise RuntimeError("Failed to save PNG file.") from e

    # Generate XML file with annotation information
    xml_root = ET.Element("annotation")
    xml_folder = ET.SubElement(xml_root, "folder")
    xml_folder.text = os.path.dirname(os.path.abspath(output_png_path))
    xml_filename = ET.SubElement(xml_root, "filename")
    xml_filename.text = os.path.basename(output_png_path)
    xml_path = ET.SubElement(xml_root, "path")
    xml_path.text = os.path.abspath(output_png_path)
    xml_source = ET.SubElement(xml_root, "source")
    xml_database = ET.SubElement(xml_source, "database")
    xml_database.text = "Unknown"
    xml_size = ET.SubElement(xml_root, "size")
    xml_width = ET.SubElement(xml_size, "width")
    xml_width.text = str(image.size[0])
    xml_height = ET.SubElement(xml_size, "height")
    xml_height.text = str(image.size[1])
    xml_depth = ET.SubElement(xml_size, "depth")
    xml_depth.text = "3"
    xml_segmented = ET.SubElement(xml_root, "segmented")
    xml_segmented.text = "0"
    for i, (target_text, annotation_text) in enumerate(target_annotation_pairs):
        xml_object = ET.SubElement(xml_root, "object")
        xml_name = ET.SubElement(xml_object, "name")
        xml_name.text = annotation_text
        xml_pose = ET.SubElement(xml_object, "pose")
        xml_pose.text = "Unspecified"
        xml_truncated = ET.SubElement(xml_object, "truncated")
        xml_truncated.text = "0"
        xml_difficult = ET.SubElement(xml_object, "difficult")
        xml_difficult.text = "0"
        bbox_start = i * len(all_bbox) // len(target_annotation_pairs)
        bbox_end = (i+1) * len(all_bbox) // len(target_annotation_pairs)
        for box in all_bbox[bbox_start:bbox_end]:
            xml_bndbox = ET.SubElement(xml_object, "bndbox")
            xml_xmin = ET.SubElement(xml_bndbox, "xmin")
            xml_xmin.text = str(box[0])
            xml_ymin = ET.SubElement(xml_bndbox, "ymin")
            xml_ymin.text = str(box[1])
            xml_xmax = ET.SubElement(xml_bndbox, "xmax")
            xml_xmax.text = str(box[2])
            xml_ymax = ET.SubElement(xml_bndbox, "ymax")
            xml_ymax.text = str(box[3])

    xml_tree = ET.ElementTree(xml_root)
    try:
        xml_tree.write(output_xml_path)
    except Exception as e:
        raise RuntimeError("Failed to save XML file.") from e



def convert_to_image(doc_file_path):
    """
    Converts a Docx file to a PIL Image.

    Args:
        doc_file_path (str): Path to the input Docx file.

    Raises:
        ValueError: If the input file does not exist.
        RuntimeError: If the conversion fails or the output file is not created.

    Returns:
        PIL Image object
    """
    # Check if input file exists
    if not os.path.isfile(doc_file_path):
        raise ValueError("Input file does not exist.")

    try:
        with tempfile.TemporaryDirectory() as path:
            # Convert docx to pdf
            convert(doc_file_path, os.path.join(path, "temp.pdf"))

            # Convert pdf to image
            images = pdf2image.convert_from_path(os.path.join(path, "temp.pdf"))
            image = images[0]  # Use first page

    except Exception as e:
        raise RuntimeError("Failed to convert Docx to image.") from e

    return image

def get_bounding_box_multi_page(target_text: str, image: Image) -> List[List[int]]:
    """
    Get bounding box coordinates for target text in an image with multiple pages.

    Args:
        target_text (str): The target text to search for.
        image (PIL Image): The input image.

    Returns:
        List of lists containing [x1, y1, x2, y2] coordinates for each bounding box.
    """
    # Initialize variables
    page_num = 0
    width, height = image.size
    page_boxes = []
    search_text = re.sub(r'[^\x00-\x7F]+', '', target_text) # Remove non-ASCII characters from text

    # Loop through each page of the image
    while True:
        try:
            page_image = image.copy()
            left = 0
            top = page_num * height
            right = width
            bottom = (page_num + 1) * height
            page_image = page_image.crop((left, top, right, bottom))

            # Perform OCR on page to get bounding boxes for target text
            pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            ocr_text = pytesseract.image_to_data(page_image, output_type=pytesseract.Output.DICT)
            boxes = []
            words = [w.lower() for w in search_text.split()]
            for i in range(len(ocr_text['text']) - len(words) + 1):
                if search_text.__contains__('\n'):
                    mod_text = ' '.join([w.lower() for w in ocr_text['text'][i:i+len(words)+1]]).lstrip()
                    mod_search_text = search_text.lower().replace('\n', '  ').strip()
                else:
                    mod_text = ' '.join([w.lower() for w in ocr_text['text'][i:i+len(words)]])
                    mod_search_text = search_text.lower()
                if mod_text == mod_search_text:
                    x = ocr_text['left'][i]
                    y = ocr_text['top'][i]
                    w = sum(ocr_text['width'][i:i+len(words)])
                    h = max(ocr_text['height'][i:i+len(words)])
                    boxes.append([x + left, y + top, x + w + left, y + h + top])

            if not boxes:
                break

            # Add bounding boxes to list for page
            page_boxes.append(boxes)

            # Move on to next page
            page_num += 1

        except IndexError:
            break

    # Combine bounding boxes for all pages into a single list
    all_boxes = []
    for boxes in page_boxes:
        for box in boxes:
            all_boxes.append(box)

    return all_boxes


def convert_docx_to_pdf(input_file_path, output_file_path):
    """
    Converts a docx file to pdf.

    Args:
        input_file_path (str): Path to the input docx file.
        output_file_path (str): Path to the output pdf file.

    Raises:
        ValueError: If the input file does not exist.
        RuntimeError: If the conversion fails or the output file is not created.

    Returns:
        None
    """
    # Check if input file exists
    if not os.path.isfile(input_file_path):
        raise ValueError("Input file does not exist.")

    try:
        # Convert docx to pdf
        convert(input_file_path, output_file_path)
    except Exception as e:
        raise RuntimeError("Failed to convert docx to pdf.") from e

    # Check if output file exists
    if not os.path.isfile(output_file_path):
        raise RuntimeError("Failed to create pdf output file.")

def add_table_to_doc(doc_path, placeholder_text, rows, cols):
    # Open the docx file
    doc = Document(doc_path)

    # Before creating the table
    print(f"Number of tables before: {len(doc.tables)}")

    # Find the placeholder
    placeholder = None
    for paragraph in doc.paragraphs:
        if placeholder_text in paragraph.text:
            placeholder = paragraph
            break

    # If placeholder is found, replace it with the table
    if placeholder:
        # Remove the placeholder paragraph
        p = placeholder._element
        p.getparent().remove(p)
        p._p = p._element = None

        # Insert an empty paragraph before the table
        placeholder.insert_paragraph_before('Tomcat')

        # Create the table
        table = doc.add_table(rows=rows+1, cols=cols)

        # Print the text of the paragraph before the table
        print(table._element.getprevious().text)

        # After creating the table
        print(f"Number of tables after: {len(doc.tables)}")

        # Set table alignment to center
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        col_widths = [Cm(2), Cm(3), Cm(4), Cm(2), Cm(2), Cm(3)]
        for i, width in enumerate(col_widths):
            table.columns[i].width = width

        # Set table header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Txn Date'
        hdr_cells[1].text = 'Description'
        hdr_cells[2].text = 'Ref No./Cheque No.'
        hdr_cells[3].text = 'Debit'
        hdr_cells[4].text = 'Credit'
        hdr_cells[5].text = 'Balance'

        # Populate table with transaction data
        fake = Faker()
        balance = 4000
        total_debit = 0
        total_credit = 0
        for i in range(1, rows):
            row_cells = table.rows[i].cells
            txn_date = fake.date_between(start_date='-1y', end_date='today').strftime('%d %b %Y')
            row_cells[0].text = txn_date
            row_cells[1].text = fake.sentence(nb_words=5, variable_nb_words=True)
            row_cells[2].text = str(uuid.UUID(fake.uuid4()).hex)[:8]
            is_debit = random.choice([True, False])
            if is_debit:
                debit = random.randint(100, 5000)
                credit = 0
            else:
                debit = 0
                credit = random.randint(100, 5000)
            balance = balance + credit - debit
            if debit > 0:
                row_cells[3].text = f"{debit:,}"
            else:
                row_cells[3].text = ''
            if credit > 0:
                row_cells[4].text = f"{credit:,}"
            else:
                row_cells[4].text = ''
            if balance >= 0:
                row_cells[5].text = f"{balance:,}"
            else:
                row_cells[5].text = ''

            total_debit += debit
            total_credit += credit

        # Add final row with total debit, total credit, and final balance
        row_cells = table.add_row().cells
        row_cells[1].text = "Balance"
        row_cells[3].text = f"{total_debit:,}"
        row_cells[4].text = f"{total_credit:,}"
        if balance >= 0:
            row_cells[5].text = f"{balance:,}"
        else:
            row_cells[5].text = ''


        # Add spacing after the table
        doc.add_paragraph(' ')

        # Save the modified docx file
        doc.save(doc_path)

        print(f"Table successfully added to {doc_path}!")
    else:
        print(f"Placeholder '{placeholder_text}' not found in {doc_path}!")

def generate_document_and_pdf(template_path, data, output_dir):
    """
    Generates a document and a pdf for each row in the given data.

    Args:
        template_path (str): Path to the document template.
        data (pd.DataFrame): Pandas DataFrame containing the data for the documents.
        output_dir (str): Path to the output directory.

    Raises:
        ValueError: If the input template docx file does not exist.
        RuntimeError: If the document or pdf generation fails or the output files are not created.

    Returns:
        None
    """
    # Check if template file exists
    if not os.path.isfile(template_path):
        raise ValueError("Template docx file does not exist.")

    for index, row in data.iterrows():
        # Define context for the document
        context = {
            'account_holder_name': row['account_holder_name'],
            'address': row['address'],
            'ifsc_code': row['ifsc_code'],
            'micr_code': row['micr_code'],
            'branch_name': row['branch_name'],
            'account_type': row['account_type'],
            'account_number': row['account_number'],
            'opening_balance': row['opening_balance'],
            'closing_balance': row['closing_balance'],
            'debit': row['debit'],
            'credit': row['credit'],
            'total_balance': row['total_balance'],
            'today_date': datetime.today().strftime("%d %b, %Y"),
        }

        # Load the document template and render it with the context
        doc = DocxTemplate(template_path)
        doc.render(context)

        # Pretty Print the current Context
        # pprint.pprint(context)
        # Define output file paths
        docx_output_path = os.path.join(output_dir, f"generated_doc_{input_yaml_data['file_name']}_{input_yaml_data['action']}_{index+1}.docx")
        pdf_output_path = os.path.join(output_dir, f"generated_doc_{input_yaml_data['file_name']}_{input_yaml_data['action']}_{index+1}.pdf")
        png_output_path = os.path.join(output_dir, f"generated_doc_{input_yaml_data['file_name']}_{input_yaml_data['action']}_{index+1}.png")
        xml_output_path = os.path.join(output_dir, f"generated_doc_{input_yaml_data['file_name']}_{input_yaml_data['action']}_{index+1}.xml")

        try:
            # Save the document as a docx file
            doc.save(docx_output_path)
            if input_yaml_data['automate_table']:
                # Add table to the document
                add_table_to_doc(docx_output_path,'Transaction Table', 10, 6) # Example values for 10 rows and 6 columns

            annotation_tuple = create_target_annotation_tuple(context_dict=context,label_list=mapping_yaml_data,target_label=input_yaml_data['target_label'])
            
            # Annotate and convert to PNG
            annotate_text(docx_output_path,annotation_tuple,png_output_path,xml_output_path)

            # Convert the docx file to pdf
            # convert_docx_to_pdf(docx_output_path, pdf_output_path)
        except Exception as e:
            raise RuntimeError("Failed to generate document and pdf.") from e

# Function to create The tuple for Annotation Label Mapping
def create_target_annotation_tuple(context_dict, label_list, target_label):
    label_map = [(x[0], x[1]) for x in label_list]
    filtered_dict = {k: v for k, v in context_dict.items() if k in target_label}
    result = []
    for key, value in label_map:
        if key in filtered_dict:
            result.append((str(filtered_dict[key]), value))
    return result


if __name__ == '__main__':
    # Define input and output directories
    input_dir = os.getcwd()
    output_doc_dir = 'output_docs'
    output_pdf_dir = 'output_pdfs'

    # Load the config input
    with open('config/input.yaml', 'r') as file:
        input_yaml_data = yaml.safe_load(file)

    # Load the config Label Mappings
    with open('config/label_mapping.yaml', 'r') as file:
        mapping_yaml_data = yaml.safe_load(file)
    
    # Define input file paths
    template_path = os.path.join(input_dir, 'docx_templates', f"{input_yaml_data['file_name']}.docx")
    data_path = os.path.join(input_dir, 'excel_sheets', f"{input_yaml_data['file_name']}_accounts.csv")

    # Read data from CSV file
    data = pd.read_csv(data_path)

    # Generate documents and pdfs
    generate_document_and_pdf(template_path, data, output_doc_dir)
    

